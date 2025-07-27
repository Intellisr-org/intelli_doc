import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from table_recognition import TableRecognition


logger = logging.getLogger(__name__)

class WordDocumentProcessor:
    def __init__(self, confidence_threshold=90.0):
        self.confidence_threshold = confidence_threshold / 100.0  # Convert percentage to decimal
    
    def create_word_document(self, image, ocr_results, layout_results, output_path):
        """Create Word document from OCR and layout results"""
        try:
            doc = Document()
            
            for i, (ocr_result, layout_result) in enumerate(zip(ocr_results, layout_results)):
                page_num = i + 1
                
                if i > 0:
                    doc.add_page_break()
                
                # Add page number as comment
                doc.add_paragraph(f"--- Page {page_num} ---", style='Heading 3')
                
                # Process based on layout if available
                if 'layout_predictions' in layout_result and layout_result['layout_predictions']:
                    self.process_layout_content(image, doc, ocr_result, layout_result)
                else:
                    # Fallback to simple text extraction with confidence filtering
                    self.process_simple_text_with_confidence(doc, ocr_result)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            return False
    
    def process_layout_content(self, image, doc, ocr_result, layout_result):
        """Process content based on layout analysis"""
        try:
            if 'error' in ocr_result:
                doc.add_paragraph(f"[Error processing page: {ocr_result['error']}]")
                return

            # get page data
            page_data = {}
            page_data['page_width'] = ocr_result.get('page_width')
            page_data['page_height'] = ocr_result.get('page_height')
            page_data['image'] = image

            # Filter OCR results by confidence threshold
            total_lines = len(ocr_result.get('text_lines', []))
            filtered_text_lines = []
            for line in ocr_result.get('text_lines', []):
                confidence = line.get('confidence', 0)
                if confidence is not None and confidence >= self.confidence_threshold:
                    filtered_text_lines.append(line)
            
            logger.info(f"Filtered {len(filtered_text_lines)}/{total_lines} text lines (confidence >= {self.confidence_threshold * 100}%)")
            
            if not filtered_text_lines:
                logger.info(f"No text lines meet the confidence threshold ({self.confidence_threshold * 100}%)")
                return
            
            # Get all layout boxes
            layout_boxes = []
            for pred in layout_result['layout_predictions']:
                if isinstance(pred, dict) and 'bboxes' in pred:
                    for bbox_data in pred['bboxes']:
                        layout_boxes.append(bbox_data)

            # get layout of undetected page-headers and page-footers
            if filtered_text_lines:
                updated_layout_boxes = self.process_undetected_headers_and_footers(filtered_text_lines, layout_boxes, page_data)

            # Find text lines that are detected by layout analysis
            layout_detected_lines = []
            for line in filtered_text_lines:
                line_bbox = line.get('bbox')
                if line_bbox:
                    # Check if this line overlaps with any layout box
                    for layout_box in updated_layout_boxes:
                        if self.bbox_overlaps(line_bbox, layout_box['bbox']):
                            layout_detected_lines.append(line)
                            break  # Found a match, no need to check other layout boxes
            
            logger.info(f"Found {len(layout_detected_lines)}/{len(filtered_text_lines)} text lines detected by layout analysis")
            
            if not layout_detected_lines:
                logger.info("No text lines are detected by layout analysis")
                return

            # Sort layout boxes by vertical position
            updated_layout_boxes.sort(key=lambda x: x['bbox'][1])  # Sort by y-coordinate
            
            # Process each layout box
            for layout_box in updated_layout_boxes:
                label = layout_box.get('label', 'Text')
                bbox = layout_box['bbox']
                
                # Find text lines within this layout box (only from filtered and layout-detected lines)
                relevant_lines = []
                relevant_lines_full = []
                for line in layout_detected_lines:
                    if line.get('bbox') and self.bbox_overlaps(line['bbox'], bbox):
                        relevant_lines.append(line['text'])
                        relevant_lines_full.append(line)
                
                if relevant_lines:
                    text = ' '.join(relevant_lines).strip()
                    if text:
                        self.add_formatted_text(doc, text, label, relevant_lines_full, layout_box, page_data,image)
                        
        except Exception as e:
            logger.error(f"Error processing layout content: {e}")
            # Fallback to simple text processing with confidence filter
            self.process_simple_text_with_confidence(doc, ocr_result)
    
    def process_simple_text_with_confidence(self, doc, ocr_result):
        """Simple text processing with confidence filtering"""
        try:
            if 'error' in ocr_result:
                doc.add_paragraph(f"[Error processing page: {ocr_result['error']}]")
                return
            
            # Extract text lines with confidence threshold
            total_lines = len(ocr_result.get('text_lines', []))
            text_lines = []
            for line in ocr_result.get('text_lines', []):
                confidence = line.get('confidence', 0)
                if confidence is not None and confidence >= self.confidence_threshold and line.get('text'):
                    text_lines.append(line['text'])
            
            logger.info(f"Using {len(text_lines)}/{total_lines} text lines (confidence >= {self.confidence_threshold * 100}%)")
            
            if text_lines:
                # Join all text
                full_text = '\n'.join(text_lines)
                doc.add_paragraph(full_text)
            else:
                logger.info(f"No text lines meet the confidence threshold ({self.confidence_threshold * 100}%)")
                
        except Exception as e:
            logger.error(f"Error processing simple text with confidence filter: {e}")
            doc.add_paragraph("[Error processing page content]")
    
    def bbox_overlaps(self, bbox1, bbox2, threshold=0.5):
        """Check if two bounding boxes overlap significantly"""
        try:
            x1, y1, x2, y2 = bbox1
            x3, y3, x4, y4 = bbox2
            
            # Calculate intersection
            left = max(x1, x3)
            top = max(y1, y3)
            right = min(x2, x4)
            bottom = min(y2, y4)
            
            if left < right and top < bottom:
                intersection_area = (right - left) * (bottom - top)
                bbox1_area = (x2 - x1) * (y2 - y1)
                overlap_ratio = intersection_area / bbox1_area if bbox1_area > 0 else 0
                return overlap_ratio > threshold
            
            return False
            
        except Exception:
            return False
    
    def add_formatted_text(self, doc, text, label, relevant_lines_full, layout_box, page_data,image):
        """Add formatted text to document based on label"""
        try:
            if label == 'Title':
                doc.add_heading(text, level=0)
            elif label == 'SectionHeader':
                doc.add_heading(text, level=1)
            elif label == 'List':
                doc.add_paragraph(text, style='List Bullet')
            elif label == 'Formula':
                p = doc.add_paragraph(text)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            elif label == 'Table':
                table_data = TableRecognition.table_recognition(relevant_lines_full, layout_box, page_data)
                result = self.add_table_to_doc(doc, table_data)
            elif label== 'TableofContents':
                table_data = TableRecognition.table_recognition(relevant_lines_full, layout_box, page_data)
                result = self.add_table_to_doc(doc, table_data)
            elif label == 'Figure':
                self.add_image_to_doc(doc, image, layout_box, page_data, text)
            elif label == 'PageHeader':
                cleaned_text = self.clean_text(text)
                p = doc.add_paragraph(cleaned_text, style='Normal')
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7)
            elif label == 'PageFooter':
                cleaned_text = self.clean_text(text)
                p = doc.add_paragraph(cleaned_text, style='Normal')
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7)
            else:
                doc.add_paragraph(text)
                
        except Exception as e:
            logger.error(f"Error formatting text: {e}")
            doc.add_paragraph(text)  # Fallback

    # Helper method for table and table-of-contents processing
    def is_within(self, bbox, polygon):
        cx = sum([pt[0] for pt in polygon]) / len(polygon)
        cy = sum([pt[1] for pt in polygon]) / len(polygon)
        return (bbox[0] <= cx <= bbox[2]) and (bbox[1] <= cy <= bbox[3])
    
    # Helper method for table and table-of-contents processing
    def clean_text(self, text):
        text = re.sub(r"<.*?>", "", text).strip()

        # Remove specific math commands
        math_commands = [
            r"\\mathbf\s*",      # \mathbf
            r"\\mathrm\s*",      # \mathrm
            r"\\mathit\s*",      # \mathit
            r"\\mathcal\s*",     # \mathcal
            r"\\mathbb\s*",      # \mathbb
            r"\\mathfrak\s*",    # \mathfrak
            r"\\mathsf\s*",      # \mathsf
            r"\\mathtt\s*",      # \mathtt
            r"\\text\s*",        # \text
            r"\\textbf\s*",      # \textbf
            r"\\textit\s*",      # \textit
        ]
        
        for cmd in math_commands:
            text = re.sub(cmd, "", text)
        return text

    # Helper method for table and table-of-contents processing
    def group_rows_by_y(self, textlines, y_thresh=95):
        """Group lines into rows based on y-center proximity."""
        rows = []
        for tl in sorted(textlines, key=lambda t: (t['polygon'][0][1])):  # sort by y
            cy = sum([pt[1] for pt in tl['polygon']]) / 4
            matched = False
            for row in rows:
                avg_cy = sum([sum([pt[1] for pt in tl_['polygon']])/4 for tl_ in row]) / len(row)
                if abs(cy - avg_cy) < y_thresh:
                    row.append(tl)
                    matched = True
                    break
            if not matched:
                rows.append([tl])
        return rows

    def add_table_to_doc(self, doc:Document, table_data):
        """Add detected table to Word document as simple grid table"""
        
        rows = table_data.get('rows', 0)
        columns = table_data.get('cols', 0)
        table_content = table_data.get('table_data', [])

        # Add empty paragraph for spacing
        doc.add_paragraph("", style='Normal')
        
        # Create table
        table = doc.add_table(rows=rows, cols=columns)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Fill table with content
        for row_idx in range(rows):
            for col_idx in range(columns):
                if row_idx < len(table_content) and col_idx < len(table_content[row_idx]):
                    cell_text = table_content[row_idx][col_idx]
                    table.cell(row_idx, col_idx).text = str(cell_text).strip()
        
        doc.add_paragraph("", style='Normal')
        
        return table


    def process_table_of_contents(self, doc: Document, relevant_lines_full, page_data):
        """Process table of contents and add it to the document."""
        try:
            # Parameters
            y_threshold = 15       # Vertical grouping tolerance
            x_threshold = 0.15     # Horizontal clustering threshold (normalized)

            # Step 1: Filter only text lines within TOC layout box
            toc_text_lines = [tl for tl in relevant_lines_full[0]['text_lines']]
            if not toc_text_lines:
                return

            # Step 2: Extract position data
            lines_with_pos = []
            for tl in toc_text_lines:
                x_coords = [pt[0] for pt in tl['polygon']]
                y_coords = [pt[1] for pt in tl['polygon']]

                left_x = min(x_coords)
                right_x = max(x_coords)
                center_x = sum(x_coords) / len(x_coords)
                center_y = sum(y_coords) / len(y_coords)

                rel_left = left_x / page_data['page_width']
                rel_center = center_x / page_data['page_width']
                rel_right = right_x / page_data['page_width']

                lines_with_pos.append({
                    'text': self.clean_text(tl['text']),
                    'cy': center_y,
                    'rel_left': rel_left,
                    'rel_center': rel_center,
                    'rel_right': rel_right
                })

            # Step 3: Group lines vertically
            lines_with_pos.sort(key=lambda x: x['cy'])
            grouped_lines = []
            current_group = []
            for line in lines_with_pos:
                cy = line['cy']
                if not current_group or abs(cy - current_group[-1]['cy']) <= y_threshold:
                    current_group.append(line)
                else:
                    grouped_lines.append(current_group)
                    current_group = [line]
            if current_group:
                grouped_lines.append(current_group)

            # Step 4: Write to document
            doc.add_paragraph("", style='Normal')
            for group in grouped_lines:
                group.sort(key=lambda x: x['rel_center'])
                x_bands = []
                for item in group:
                    rel_x = item['rel_center']
                    assigned = False
                    for band in x_bands:
                        if abs(rel_x - band['center']) <= x_threshold:
                            band['items'].append(item)
                            band['center'] = sum(x['rel_center'] for x in band['items']) / len(band['items'])
                            assigned = True
                            break
                    if not assigned:
                        x_bands.append({'center': rel_x, 'items': [item]})
                x_bands.sort(key=lambda b: b['center'])
                left_text = ""
                middle_text = ""
                right_text = ""
                for band in x_bands:
                    texts = [x['text'] for x in band['items']]
                    band_text = " ".join(texts).strip()
                    sample_item = band['items'][0]
                    rel_left = sample_item['rel_left']
                    rel_right = sample_item['rel_right']
                    if rel_left < 0.1:
                        left_text += band_text + " "
                    elif rel_right > 0.75:
                        right_text += band_text
                    else:
                        middle_text += band_text + " "
                dot_space = " " * 30 + "." * 10 + " " if right_text else " "
                line = f"{left_text.strip():<25}{middle_text.strip()}{dot_space}{right_text.strip():>8}"
                doc.add_paragraph(line, style='No Spacing')
            doc.add_paragraph("", style='Normal')
        except Exception as e:
            print("Error: ", e)

    def process_undetected_headers_and_footers(self, filtered_text_lines, layout_boxes, page_data):
        """Handle undetected layout boxes (e.g., headers and footers)"""
        try:
            top_20 = []
            bottom_20 = []

            # Initialize new_layout_boxes with a copy of the current layout_boxes.

            new_layout_boxes = list(layout_boxes)

            for tl in filtered_text_lines:
                result = self.is_in_top_or_bottom_20(tl['polygon'], page_data['page_height'])
                if result['in_top_20']:
                    top_20.append(tl)
                if result['in_bottom_20']:
                    bottom_20.append(tl)

            # Check if any polygon in top 20 is within any existing layout box
            top_20_not_within_layout = []
            bottom_20_not_within_layout = []
            for tl in top_20:
                # Check against the *original* layout_boxes to determine if it's "not within layout"
                if not any(self.is_within(layout_box['bbox'], tl['polygon']) for layout_box in layout_boxes):
                    top_20_not_within_layout.append(tl)
            for tl in bottom_20:
                # Check against the *original* layout_boxes
                if not any(self.is_within(layout_box['bbox'], tl['polygon']) for layout_box in layout_boxes):
                    bottom_20_not_within_layout.append(tl)

            # Convert undetected text boxes into new layout boxes and add to new_layout_boxes
            for tl in top_20_not_within_layout:
                # self.convert_textbox_to_layoutbox modifies new_layout_boxes in place
                # and returns the modified list.
                new_layout_boxes = self.convert_textbox_to_layoutbox(tl, new_layout_boxes, 'Page-header')
            for tl in bottom_20_not_within_layout:
                # self.convert_textbox_to_layoutbox modifies new_layout_boxes in place
                new_layout_boxes = self.convert_textbox_to_layoutbox(tl, new_layout_boxes, 'PageFooter')

            # Now, iterate through the top 20 text boxes to identify potential page numbers
            # and update their labels in the 'new_layout_boxes' list.
            for tl in top_20:
                text = tl.get('text', '').strip()
                is_bold = tl.get('is_bold', False)
                if text.isdigit() and not is_bold:
                    # Find this specific text box's corresponding layout box in new_layout_boxes
                    # and update its label to 'PageHeader'.
                    for layout_box in new_layout_boxes:
                        # Compare the bounding boxes to find the matching layout box.
                        if self.are_bboxes_approximately_equal(layout_box['bbox'], tl['bbox']):
                            layout_box['label'] = 'Page-header'
                            print(f"Updated layout box label to 'Page-header' for bbox: {tl['bbox']}")
                            break # Once found and updated, no need to check further in new_layout_boxes

            # Iterate through the bottom 20 text boxes to identify potential page numbers
            # and update their labels in the 'new_layout_boxes' list.
            for tl in bottom_20:
                text = tl.get('text', '').strip()
                is_bold = tl.get('is_bold', False)
                if text.isdigit() and not is_bold:
                    # Find this specific text box's corresponding layout box in new_layout_boxes
                    # and update its label to 'PageFooter'.
                    for layout_box in new_layout_boxes:
                        # Compare the bounding boxes to find the matching layout box.
                        if self.are_bboxes_approximately_equal(layout_box['bbox'], tl['bbox']):
                            layout_box['label'] = 'PageFooter'
                            print(f"Updated layout box label to 'PageFooter' for bbox: {tl['bbox']}")
                            break # Once found and updated, no need to check further in new_layout_boxes

            return new_layout_boxes

        except Exception as e:
            print("Error process_undetected_headers_and_footers: ", e)

    def are_bboxes_approximately_equal(self, bbox1, bbox2, epsilon=1.0):
        """
        Checks if two bounding boxes are approximately equal within a given epsilon.
        BBox format: [x1, y1, x2, y2]
        """
        if len(bbox1) != 4 or len(bbox2) != 4:
            return False # Bounding boxes must have 4 coordinates

        # Compare each coordinate with the given tolerance
        for i in range(4):
            if abs(bbox1[i] - bbox2[i]) > epsilon:
                return False # If any coordinate differs by more than epsilon, they are not approximately equal
        return True # All coordinates are within the tolerance
    
    # Helper method for considering undetected layout boxes
    def is_in_top_or_bottom_20(self, polygon, image_height):
        y_coords = [pt[1] for pt in polygon]
        min_y = min(y_coords)
        max_y = max(y_coords)

        top_20_y = image_height * 0.15
        bottom_20_y = image_height * 0.85

        return {
            'in_top_20': min_y <= top_20_y,
            'in_bottom_20': max_y >= bottom_20_y
        }

    # Helper method for considering undetected layout boxes
    def convert_textbox_to_layoutbox(self, tl, layout_boxes, label):
        text_bbox = tl['bbox']
        polygon = tl['polygon']

        # Determine the type of the layout box to create.
        # If layout_boxes is empty, assume a dictionary structure as per your example.
        if layout_boxes:
            # Use the type of an existing box to maintain consistency
            existing_box_type = type(layout_boxes[0])
        else:
            # If layout_boxes is empty, default to dict, assuming the structure is a dictionary
            existing_box_type = dict

        new_layout_box = existing_box_type(
            bbox=text_bbox,
            polygon=polygon,
            confidence=0.8,  # You can use the text confidence or set your own
            label=label,
        )
        layout_boxes.append(new_layout_box)

        return layout_boxes

    def add_image_to_doc(self, doc, image, layout_box, page_data ,text):
        """Add image to document based on layout box"""
        try:
            # Get image dimensions
            image_width, image_height = image.size
            
            layout_box_width = layout_box['bbox'][2] - layout_box['bbox'][0]
            layout_box_height = layout_box['bbox'][3] - layout_box['bbox'][1]

            # Calculate scaling factors
            scale_x = layout_box_width / image_width
            scale_y = layout_box_height / image_height

            image_width = layout_box_width * scale_x
            image_height = layout_box_height * scale_y
            
            # Create a new image with the calculated dimensions
            new_image = image.resize((int(image_width), int(image_height)))
            
            # Add the image to the document
            doc.add_picture(new_image, width=Inches(layout_box_width))
            doc.add_paragraph(f"[Figure: {text}]")
            return True
        except Exception as e:
            print("Error add_image_to_doc: ", e)
            return False